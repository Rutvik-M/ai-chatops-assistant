# create_kb_files.py
import os
from docx import Document
from fpdf import FPDF

# Create directories
os.makedirs("knowledge_base/general", exist_ok=True)
os.makedirs("knowledge_base/hr", exist_ok=True)
os.makedirs("knowledge_base/finance", exist_ok=True)
os.makedirs("knowledge_base/it_support", exist_ok=True)
os.makedirs("knowledge_base/engineering", exist_ok=True)

print("📁 Creating knowledge base files...")

# ========== GENERAL FOLDER ==========

# 1. company_policy.txt
with open("knowledge_base/general/company_policy.txt", "w") as f:
    f.write("""COMPANY POLICIES AND GUIDELINES

Last Updated: November 2025

1. WORK HOURS
- Standard work hours: 9:00 AM - 6:00 PM (Monday to Friday)
- Flexible hours available with manager approval
- Core hours (mandatory): 10:00 AM - 3:00 PM

2. DRESS CODE
- Business casual attire required
- Casual Fridays: Jeans and polo shirts allowed
- Client-facing roles: Business formal attire

3. ATTENDANCE POLICY
- Punctuality is expected
- Late arrivals (>15 mins) must be reported to manager
- 3 unexcused late arrivals = written warning

4. REMOTE WORK POLICY
- Hybrid work model available (3 days office, 2 days remote)
- Full remote work requires VP approval
- Remote employees must maintain core hours availability

5. CODE OF CONDUCT
- Respect all colleagues regardless of position
- Zero tolerance for harassment or discrimination
- Professional communication expected at all times

6. CONFIDENTIALITY
- All company information is confidential
- NDAs must be signed by all employees
- Data breaches must be reported immediately

7. EQUIPMENT USAGE
- Company equipment for business use only
- Personal use allowed during breaks (reasonable limits)
- Lost/damaged equipment must be reported within 24 hours

For questions, contact HR at hr@company.com
""")

# 2. Remote_Work_Guidelines.md
with open("knowledge_base/general/Remote_Work_Guidelines.md", "w") as f:
    f.write("""# Remote Work Guidelines

## Eligibility
- All full-time employees after 90 days of employment
- Manager approval required
- Performance review score must be "Meets Expectations" or higher

## Equipment Provided
- Laptop computer (Dell or HP)
- External monitor (24" or 27")
- Keyboard and mouse
- Headset with microphone
- Docking station

## Technical Requirements
- Minimum internet speed: 50 Mbps download, 10 Mbps upload
- Dedicated workspace with door (privacy required)
- Backup internet connection (mobile hotspot acceptable)
- Updated operating system and security software

## Work Schedule
- Must be available during core hours (10 AM - 3 PM local time)
- Flexible start/end times outside core hours
- Same total hours as office schedule (40 hours/week)

## Communication Expectations
- Respond to Slack/Teams messages within 30 minutes during work hours
- Camera on for all team meetings
- Daily standup attendance mandatory
- Weekly 1-on-1 with manager

## Security Requirements
- Use company VPN at all times
- Lock computer when stepping away
- Secure home Wi-Fi with WPA2/WPA3 encryption
- No public Wi-Fi for work activities

## Performance Monitoring
- Same performance standards as office workers
- Weekly deliverables tracking
- Quarterly remote work review with manager

Contact IT Support for technical issues: it@company.com
""")

# ========== HR FOLDER ==========

# 3. pto_policy.txt
with open("knowledge_base/hr/pto_policy.txt", "w") as f:
    f.write("""PAID TIME OFF (PTO) POLICY

Effective Date: January 1, 2025

ACCRUAL RATES:
- 0-2 years of service: 15 days per year (1.25 days per month)
- 3-5 years of service: 20 days per year (1.67 days per month)
- 6+ years of service: 25 days per year (2.08 days per month)

REQUESTING PTO:
1. Submit request via HR portal at least 2 weeks in advance
2. Manager approval required within 3 business days
3. Emergency PTO: Notify manager immediately via phone/text

BLACKOUT PERIODS:
- December 15-31 (limited approvals, first-come basis)
- End of fiscal quarters (limited to 50% of team)

CARRYOVER POLICY:
- Maximum 5 unused days can be carried to next year
- Excess days expire on December 31st
- No payout for unused PTO upon separation (unless state law requires)

SICK LEAVE:
- 10 sick days per year (separate from PTO)
- Medical documentation required for 3+ consecutive days
- Short-term disability available after sick days exhausted

HOLIDAYS (12 days, company closed):
- New Year's Day
- Martin Luther King Jr. Day
- Presidents' Day
- Memorial Day
- Independence Day
- Labor Day
- Thanksgiving (2 days)
- Christmas (3 days)

PARENTAL LEAVE:
- Primary caregiver: 12 weeks paid
- Secondary caregiver: 6 weeks paid
- Available to all parents (birth, adoption, foster)

BEREAVEMENT LEAVE:
- Immediate family: 5 days paid
- Extended family: 3 days paid
- Additional time available as unpaid leave

For PTO balance inquiries, contact: hr@company.com
""")

# 4. performance_reviews.docx
doc = Document()
doc.add_heading('Performance Review Process', 0)

doc.add_heading('Review Schedule', 1)
doc.add_paragraph('Annual Performance Reviews: Conducted every December')
doc.add_paragraph('Mid-Year Check-ins: June (informal)')
doc.add_paragraph('Quarterly 1-on-1s: Required with direct manager')

doc.add_heading('Rating Scale', 1)
doc.add_paragraph('5 - Outstanding: Consistently exceeds all expectations')
doc.add_paragraph('4 - Exceeds Expectations: Regularly surpasses goals')
doc.add_paragraph('3 - Meets Expectations: Achieves all required objectives')
doc.add_paragraph('2 - Below Expectations: Improvement needed in key areas')
doc.add_paragraph('1 - Unsatisfactory: Does not meet minimum requirements')

doc.add_heading('Review Components', 1)
doc.add_paragraph('1. Goal Achievement (40%)')
doc.add_paragraph('2. Core Competencies (30%)')
doc.add_paragraph('3. Leadership/Teamwork (20%)')
doc.add_paragraph('4. Innovation/Initiative (10%)')

doc.add_heading('Compensation Review', 1)
doc.add_paragraph('Merit increases based on performance rating:')
doc.add_paragraph('- Rating 5: 8-10% increase')
doc.add_paragraph('- Rating 4: 5-7% increase')
doc.add_paragraph('- Rating 3: 2-4% increase')
doc.add_paragraph('- Rating 2: 0% increase (performance improvement plan)')
doc.add_paragraph('- Rating 1: 0% increase (potential termination)')

doc.add_heading('Bonus Eligibility', 1)
doc.add_paragraph('Annual bonus pool distributed based on:')
doc.add_paragraph('- Individual performance rating (60%)')
doc.add_paragraph('- Company performance (40%)')
doc.add_paragraph('Target bonus: 10-20% of base salary')

doc.save('knowledge_base/hr/performance_reviews.docx')

# ========== FINANCE FOLDER ==========

# 5. expense_policy.txt
with open("knowledge_base/finance/expense_policy.txt", "w") as f:
    f.write("""EMPLOYEE EXPENSE REIMBURSEMENT POLICY

APPROVAL LIMITS (Per Transaction):
- Individual Contributors: Up to $500
- Team Leads: Up to $2,000
- Managers: Up to $5,000
- Directors: Up to $15,000
- VP and above: Up to $50,000

REIMBURSABLE EXPENSES:

TRAVEL:
- Flights: Economy class for domestic, premium economy for international (>6 hours)
- Hotels: Up to $200/night (major cities), $150/night (other locations)
- Meals: $75/day ($15 breakfast, $25 lunch, $35 dinner)
- Ground transportation: Taxis, rideshare, rental cars, parking

BUSINESS MEALS:
- Client meals: Fully reimbursable with business purpose
- Team meals: Up to $50/person, manager approval required
- Alcohol: 2 drinks maximum per person

OFFICE SUPPLIES:
- Standard supplies: Pre-approved up to $100/month
- Specialized equipment: Requires manager approval

PROFESSIONAL DEVELOPMENT:
- Conferences: Up to $3,000/year (registration + travel)
- Training courses: Up to $2,000/year
- Books/subscriptions: Up to $500/year

SUBMISSION PROCESS:
1. Submit expenses within 30 days of transaction
2. Include itemized receipts (required for $25+)
3. Use expense management system (Expensify)
4. Provide business purpose for all expenses

REIMBURSEMENT TIMELINE:
- Standard: 10 business days after approval
- Direct deposit to bank account on file

NON-REIMBURSABLE:
- Personal expenses
- Spouse/family travel (unless approved)
- Excess baggage fees
- Minibar charges
- Late fees/penalties
- First class upgrades

CORPORATE CREDIT CARDS:
- Available to managers and above
- Monthly reconciliation required
- Personal charges strictly prohibited

Questions? Contact: finance@company.com
""")

# 6. budget_approval.pdf
pdf = FPDF()
pdf.add_page()
pdf.set_font("Arial", "B", 16)
pdf.cell(0, 10, "BUDGET APPROVAL WORKFLOW", 0, 1, "C")

pdf.set_font("Arial", "", 12)
pdf.ln(5)
pdf.multi_cell(0, 6, "QUARTERLY BUDGET APPROVAL PROCESS")
pdf.ln(3)

pdf.multi_cell(0, 6, "Step 1: Department Budget Request")
pdf.multi_cell(0, 6, "- Submit by 15th of last month of quarter")
pdf.multi_cell(0, 6, "- Include detailed line items and justification")
pdf.multi_cell(0, 6, "- Use standardized budget template")
pdf.ln(2)

pdf.multi_cell(0, 6, "Step 2: Finance Review")
pdf.multi_cell(0, 6, "- Finance team reviews within 5 business days")
pdf.multi_cell(0, 6, "- May request additional information")
pdf.multi_cell(0, 6, "- Checks against company financial targets")
pdf.ln(2)

pdf.multi_cell(0, 6, "Step 3: Executive Approval")
pdf.multi_cell(0, 6, "- Budgets >$100K require CFO approval")
pdf.multi_cell(0, 6, "- Budgets >$500K require CEO approval")
pdf.multi_cell(0, 6, "- Board approval for budgets >$2M")
pdf.ln(2)

pdf.multi_cell(0, 6, "Step 4: Implementation")
pdf.multi_cell(0, 6, "- Approved budgets active on quarter start date")
pdf.multi_cell(0, 6, "- Purchase orders can be issued")
pdf.multi_cell(0, 6, "- Monthly variance reporting required")
pdf.ln(2)

pdf.multi_cell(0, 6, "EMERGENCY BUDGET REQUESTS:")
pdf.multi_cell(0, 6, "- Flag as 'URGENT' in subject line")
pdf.multi_cell(0, 6, "- Requires VP approval within 48 hours")
pdf.multi_cell(0, 6, "- Detailed business justification mandatory")
pdf.ln(2)

pdf.multi_cell(0, 6, "Contact: finance@company.com | Ext: 5100")

pdf.output("knowledge_base/finance/budget_approval.pdf")

# ========== IT SUPPORT FOLDER ==========

# 7. it_guidelines.txt
with open("knowledge_base/it_support/it_guidelines.txt", "w") as f:
    f.write("""IT SUPPORT GUIDELINES AND TROUBLESHOOTING

COMMON ISSUES AND SOLUTIONS:

1. VPN CONNECTION ISSUES
Problem: Cannot connect to company VPN
Solutions:
- Verify internet connection is active
- Restart VPN client application
- Check VPN credentials (username: firstname.lastname)
- Try alternate VPN server from dropdown
- Contact IT if issue persists: it@company.com

VPN Server Addresses:
- Primary: vpn.company.com
- Backup: vpn2.company.com
- Asia region: vpn-asia.company.com

2. PASSWORD RESET
- Self-service: portal.company.com/reset
- Phone: Call IT helpdesk at ext. 5000
- Email locked out: Use personal email to request reset
- Expires every 90 days (notification sent 7 days prior)

Password Requirements:
- Minimum 12 characters
- At least 1 uppercase, 1 lowercase, 1 number, 1 special character
- Cannot reuse last 5 passwords

3. EMAIL ISSUES
- Outlook not syncing: Check VPN connection first
- Can't send emails: Verify email size <25MB
- Spam folder: Check weekly for false positives
- Out of storage: Archive old emails to local PST

4. WIFI ACCESS
- Office WiFi: CompanyNet (password changes monthly)
- Guest WiFi: CompanyGuest (password: Welcome2025)
- Conference rooms: Hardwired ethernet available

5. SOFTWARE INSTALLATION
- Pre-approved software: Available in Software Center
- Custom software: Submit ticket with business justification
- Approval time: 3-5 business days
- Licenses: Managed by IT, do not purchase independently

6. HARDWARE ISSUES
Laptop problems:
- Won't turn on: Hold power button 30 seconds, try again
- Slow performance: Restart weekly, close unused applications
- Overheating: Ensure vents are clear, use laptop cooling pad

Peripheral issues:
- Monitor not detected: Check cable connection, restart computer
- Keyboard/mouse: Try different USB port, replace batteries

7. PRINTER ACCESS
- Network printers: Auto-configured on domain-joined computers
- Print queue stuck: Restart print spooler service
- Mobile printing: Use PrinterOn app with employee ID

8. SECURITY BEST PRACTICES
- Lock computer when away (Windows+L)
- Never share passwords
- Report phishing emails to security@company.com
- Use password manager (LastPass provided)
- Enable MFA on all accounts

SUPPORT CONTACT METHODS:
- Phone: Ext. 5000 (8 AM - 6 PM)
- Email: it@company.com
- Ticket portal: support.company.com
- Emergency after hours: +1-555-IT-URGENT

SLA Response Times:
- Critical (system down): 1 hour
- High (work impaired): 4 hours
- Medium (inconvenience): 1 business day
- Low (question): 3 business days
""")

# 8. network_access.md
# *** FIX APPLIED HERE: Using r""" for the string literal ***
with open("knowledge_base/it_support/network_access.md", "w") as f:
    f.write(r"""# Network Access and Security Guidelines

## Network Credentials

### Active Directory Login
- Username format: `firstname.lastname`
- Domain: `COMPANY`
- Full login: `COMPANY\firstname.lastname`

### WiFi Networks

#### Primary Office Network: **CompanyNet**
- SSID: `CompanyNet`
- Security: WPA2-Enterprise
- Authentication: Use AD credentials
- Certificate: Auto-installed on company devices

#### Guest Network: **CompanyGuest**
- SSID: `CompanyGuest`
- Password: `Welcome2025` (changes quarterly)
- No access to internal resources
- Internet only, speed limited to 10 Mbps

#### Executive Network: **CompanyExec**
- SSID: `CompanyExec-5G`
- Password: Contact IT for access
- Reserved for C-level and VPs
- Priority QoS, no bandwidth limits

## VPN Access

### Cisco AnyConnect VPN
- Download: `vpn.company.com/download`
- Server: `vpn.company.com`
- Group: `RemoteAccess`

### Connection Steps:
1. Launch Cisco AnyConnect
2. Enter server: `vpn.company.com`
3. Select group: RemoteAccess
4. Enter AD username and password
5. Enter MFA code from authenticator app

### Troubleshooting VPN:
- Clear saved credentials if login fails
- Reinstall client if connection drops repeatedly
- Use backup server: `vpn2.company.com`

## File Server Access

### Shared Drives
- Department drives: \\fileserver\departments\[dept_name]
- Personal drives: \\fileserver\users\[username]
- Project drives: \\fileserver\projects\[project_code]

### Mapped Drive Letters:
- H: Drive - Personal home folder
- S: Drive - Shared departmental folder
- P: Drive - Current project folder

## Remote Desktop

### Accessing Work Computer Remotely
1. Connect to VPN first
2. Open Remote Desktop Connection
3. Enter computer name: `[username]-PC.company.com`
4. Login with AD credentials

### Requirements:
- Computer must be powered on
- Connected to network
- RDP enabled by IT

## Security Policies

### Multi-Factor Authentication (MFA)
- Required for: VPN, email, file servers
- Recommended app: Microsoft Authenticator
- Backup codes: Save in secure location

### Password Policy
- Change every 90 days
- 12 characters minimum
- Complexity required
- No password reuse (last 5)

### Device Management
- All devices must be encrypted (BitLocker)
- Antivirus required (Symantec Endpoint Protection)
- Automatic updates enabled
- Screen lock after 10 minutes idle

## Firewall Rules

### Allowed Ports:
- 80, 443 (HTTP/HTTPS)
- 22 (SSH)
- 3389 (RDP)
- 445 (SMB file sharing)

### Blocked by Default:
- Peer-to-peer file sharing
- Torrent protocols
- Cryptocurrency mining
- Gaming servers

## Cloud Services

### Approved Services:
- Microsoft 365 (email, OneDrive, Teams)
- Google Workspace (limited to contractors)
- Zoom (video conferencing)
- Slack (team communication)
- GitHub (code repositories)

### Unapproved Services:
- Personal Dropbox, Google Drive
- WeTransfer
- WhatsApp for business communication
- Personal email for work files

## Contact IT Support
- Email: it@company.com
- Phone: Ext. 5000
- Portal: support.company.com
- Emergency: +1-555-IT-URGENT
""")

# ========== ENGINEERING FOLDER ==========

# 9. deployment_guide.txt
with open("knowledge_base/engineering/deployment_guide.txt", "w") as f:
    f.write("""CODE DEPLOYMENT PROCEDURES

ENVIRONMENTS:

1. DEVELOPMENT (dev.company.com)
- Purpose: Active development and testing
- Deployment: Automatic on git push to 'dev' branch
- Access: All engineers
- Data: Synthetic test data only

2. STAGING (staging.company.com)
- Purpose: Pre-production testing and QA
- Deployment: Manual, requires PR approval
- Access: Engineers, QA team, Product managers
- Data: Anonymized production data copy

3. PRODUCTION (company.com)
- Purpose: Live customer-facing environment
- Deployment: Manual, requires multiple approvals
- Access: Senior engineers, DevOps team
- Data: Real customer data (GDPR compliant)

DEPLOYMENT PROCESS:

STEP 1: CODE REVIEW
- Create pull request (PR) on GitHub
- Minimum 2 code reviews required
- All CI/CD checks must pass
- Security scan must show no critical issues

STEP 2: STAGING DEPLOYMENT
- Merge PR to 'staging' branch
- Automatic deployment via Jenkins
- Run smoke tests and integration tests
- QA team performs manual testing

STEP 3: PRODUCTION DEPLOYMENT
- Requires approval from:
  * Tech Lead
  * Engineering Manager
  * DevOps Engineer
- Scheduled during maintenance window (Tuesdays 2-4 AM EST)
- Zero-downtime deployment using blue-green strategy

STEP 4: POST-DEPLOYMENT
- Monitor logs for errors (first 30 minutes critical)
- Check application performance metrics
- Verify customer-facing features
- Rollback plan ready if issues detected

ROLLBACK PROCEDURE:
1. Alert #engineering-alerts Slack channel
2. DevOps initiates rollback (previous version restored)
3. Post-mortem scheduled within 24 hours
4. Root cause analysis documented

DEPLOYMENT SCHEDULE:
- Deployments allowed: Mon-Thu (no Fridays!)
- Maintenance window: Tuesdays 2-4 AM EST
- Emergency hotfixes: Anytime with VP approval

BRANCHING STRATEGY:
- main: Production code
- staging: Pre-production testing
- dev: Active development
- feature/*: New features
- bugfix/*: Bug fixes
- hotfix/*: Emergency production fixes

CI/CD PIPELINE:
1. Lint and code style checks
2. Unit tests (must have >80% coverage)
3. Integration tests
4. Security vulnerability scan
5. Build Docker image
6. Deploy to environment

MONITORING:
- Application: Datadog APM
- Logs: ELK Stack (Elasticsearch, Logstash, Kibana)
- Alerts: PagerDuty
- Uptime: StatusPage.io

REQUIRED DOCUMENTATION:
- Deployment checklist completed
- Change management ticket filed
- Database migration plan (if applicable)
- Feature flags configured (if applicable)

CONTACT:
- DevOps Team: devops@company.com
- Engineering Manager: engineering@company.com
- Emergency: Slack #engineering-urgent
""")

# 10. coding_standards.md
with open("knowledge_base/engineering/coding_standards.md", "w") as f:
    f.write("""# Coding Standards and Best Practices

## General Principles

### Code Quality
- Write self-documenting code with clear variable/function names
- Keep functions small and focused (single responsibility principle)
- Avoid deep nesting (max 3 levels)
- No magic numbers - use named constants

### Comments
- Write comments for "why", not "what"
- Document complex algorithms
- Keep comments up-to-date with code changes
- Remove commented-out code before committing

## Language-Specific Standards

### Python
- **Follow PEP 8:** Use 4 spaces for indentation.
- **Naming Conventions:**
    - Modules: `lowercase_with_underscores`
    - Packages: `lowercase`
    - Classes: `CapWords` (CamelCase)
    - Functions/Variables: `lowercase_with_underscores`
    - Constants: `ALL_CAPS_WITH_UNDERSCORES`
- **Docstrings:** Use docstrings for all modules, classes, and public functions (PEP 257 format).
- **Imports:**
    - Place imports at the top of the file.
    - Group imports in this order: standard library, third-party, local application/library.
    - Separate groups with a blank line.
- **Line Length:** Limit lines to a maximum of **79 characters**.
- **Whitespace:**
    - Use blank lines sparingly to separate logical sections.
    - Avoid extraneous whitespace in expressions (e.g., `a = b + c` is preferred over `a = b + c`).
- **Function/Method Arguments:** Explicitly pass `self` as the first argument to instance methods. Use `cls` for class methods.

### JavaScript
- Use **ES6+** features.
- Prefer **const** and **let** over `var`.
- Use **semicolons** consistently.
- Use **CamelCase** for variables and functions.
- Use **PascalCase** for classes and components.
- Use **JSDoc** for documentation.

### SQL
- Use **ALL CAPS** for SQL keywords (e.g., SELECT, FROM, WHERE).
- Use **lowercase_with_underscores** for table and column names.
- Always use **prepared statements** to prevent SQL injection.
- Alias columns and tables to improve readability.
""")

print("✅ Knowledge base files created successfully.")